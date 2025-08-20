from PIL import Image
import wave
import cv2
import numpy as np
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
import moviepy.editor as mp
import io

class SteganographyProcessor:
    # Image steganography (LSB method)
    def hide_in_image(self, input_path, data, output_path):
        img = Image.open(input_path)
        binary_data = self._text_to_binary(data) + '1111111111111110'  # EOF marker
        
        if len(binary_data) > img.width * img.height * 3:
            raise ValueError("Data too large for image")
        
        pixels = list(img.getdata())
        data_index = 0
        
        for i in range(len(pixels)):
            pixel = list(pixels[i])
            for j in range(3):  # RGB channels
                if data_index < len(binary_data):
                    pixel[j] = pixel[j] & ~1 | int(binary_data[data_index])
                    data_index += 1
            pixels[i] = tuple(pixel)
        
        stego_image = Image.new(img.mode, img.size)
        stego_image.putdata(pixels)
        stego_image.save(output_path)
    
    def extract_from_image(self, input_path):
        img = Image.open(input_path)
        binary_data = []
        
        for pixel in img.getdata():
            for value in pixel[:3]:  # Only RGB
                binary_data.append(str(value & 1))
        
        binary_data = ''.join(binary_data)
        # Split by EOF marker
        eof_index = binary_data.find('1111111111111110')
        if eof_index == -1:
            raise ValueError("No hidden data found")
        
        binary_data = binary_data[:eof_index]
        return self._binary_to_text(binary_data)
    
    # Audio steganography (LSB method)
    def hide_in_audio(self, input_path, data, output_path):
        audio = wave.open(input_path, mode='rb')
        frames = bytearray(list(audio.readframes(audio.getnframes())))
        binary_data = self._text_to_binary(data) + '1111111111111110'
        
        if len(binary_data) > len(frames):
            raise ValueError("Data too large for audio")
        
        for i in range(len(binary_data)):
            frames[i] = (frames[i] & 254) | int(binary_data[i])
        
        with wave.open(output_path, 'wb') as fd:
            fd.setparams(audio.getparams())
            fd.writeframes(frames)
        audio.close()
    
    def extract_from_audio(self, input_path):
        audio = wave.open(input_path, mode='rb')
        frames = audio.readframes(audio.getnframes())
        audio.close()
        
        binary_data = []
        for frame in frames[:len(frames)]:
            binary_data.append(str(frame & 1))
        
        binary_data = ''.join(binary_data)
        eof_index = binary_data.find('1111111111111110')
        if eof_index == -1:
            raise ValueError("No hidden data found")
        
        return self._binary_to_text(binary_data[:eof_index])
    
    # Document steganography (for text files and docx)
    def hide_in_document(self, input_path, data, output_path):
        if input_path.endswith('.txt'):
            with open(input_path, 'r') as f:
                content = f.read()
            with open(output_path, 'w') as f:
                f.write(content)
                f.write(f"\n<!--{data}-->")  # Hidden in comment
        elif input_path.endswith('.docx'):
            doc = Document(input_path)
            doc.add_paragraph(data, style='CommentText')
            doc.save(output_path)
        elif input_path.endswith('.pdf'):
            reader = PdfReader(input_path)
            writer = PdfWriter()
            
            for page in reader.pages:
                writer.add_page(page)
            
            writer.add_metadata({"/HiddenData": data})
            with open(output_path, "wb") as f:
                writer.write(f)
    
    def extract_from_document(self, input_path):
        if input_path.endswith('.txt'):
            with open(input_path, 'r') as f:
                content = f.read()
            if '<!--' in content:
                return content.split('<!--')[-1].split('-->')[0]
            return ""
        elif input_path.endswith('.docx'):
            doc = Document(input_path)
            for para in doc.paragraphs:
                if para.style.name == 'CommentText':
                    return para.text
            return ""
        elif input_path.endswith('.pdf'):
            reader = PdfReader(input_path)
            return reader.metadata.get('/HiddenData', '')
    
    # Video steganography (simplified version using OpenCV)
    def hide_in_video(self, input_path, data, output_path):
        """Simplified video steganography implementation using OpenCV"""
        try:
            # Use OpenCV instead of MoviePy for simplicity
            cap = cv2.VideoCapture(input_path)
            fps = int(cap.get(cv2.CAP_PROP_FPS))
            width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            
            # Initialize video writer
            fourcc = cv2.VideoWriter_fourcc(*'mp4v')
            out = cv2.VideoWriter(output_path, fourcc, fps, (width, height))
            
            binary_data = self._text_to_binary(data) + '1111111111111110'
            data_index = 0
            frame_count = 0
            
            while True:
                ret, frame = cap.read()
                if not ret:
                    break
                
                # Only process every 5th frame to make it faster
                if frame_count % 5 == 0 and data_index < len(binary_data):
                    # Process frame to hide data
                    for i in range(height):
                        for j in range(width):
                            for channel in range(3):  # RGB channels
                                if data_index < len(binary_data):
                                    frame[i][j][channel] = (frame[i][j][channel] & 0xFE) | int(binary_data[data_index])
                                    data_index += 1
                                else:
                                    break
                            if data_index >= len(binary_data):
                                break
                        if data_index >= len(binary_data):
                            break
                
                out.write(frame)
                frame_count += 1
            
            cap.release()
            out.release()
            
        except Exception as e:
            raise Exception(f"Video processing error: {str(e)}")
    
    def extract_from_video(self, input_path):
        """Extract data from video using OpenCV"""
        try:
            cap = cv2.VideoCapture(input_path)
            binary_data = []
            eof_marker = '1111111111111110'
            frame_count = 0
            
            while True:
                ret, frame = cap.read()
                if not ret:
                    break
                
                # Only check every 5th frame
                if frame_count % 5 == 0:
                    height, width = frame.shape[:2]
                    for i in range(height):
                        for j in range(width):
                            for channel in range(3):  # RGB channels
                                binary_data.append(str(frame[i][j][channel] & 1))
                    
                    # Check if we've found the EOF marker
                    current_data = ''.join(binary_data)
                    if eof_marker in current_data:
                        extracted = current_data.split(eof_marker)[0]
                        cap.release()
                        return self._binary_to_text(extracted)
                
                frame_count += 1
            
            cap.release()
            return ""
            
        except Exception as e:
            raise Exception(f"Video extraction error: {str(e)}")
    
    # Helper methods
    def _text_to_binary(self, text):
        return ''.join(format(ord(char), '08b') for char in text)
    
    def _binary_to_text(self, binary):
        return ''.join(chr(int(binary[i:i+8], 2)) for i in range(0, len(binary), 8))